"""Build a downloadable .docx of a generated business plan.

Renders text + tables natively (no embedded charts — those are HTML-only;
users who need charts can use the HTML download).
"""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ---------- shared label dictionaries (mirror standalonePlanHtml.js) ----------
L = {
    "ru": {
        "title": "Бизнес-план",
        "feasibility": "Оценка реалистичности проекта",
        "verdicts": {"high": "Высокая", "medium": "Средняя", "low": "Низкая"},
        "executive": "Резюме",
        "market": "Контекст рынка",
        "metrics": "Финансовые показатели",
        "monthlyRevenue": "Выручка / мес.",
        "monthlyCosts": "Расходы / мес.",
        "monthlyProfit": "Прибыль / мес.",
        "breakeven": "Точка безубыточности",
        "grossMargin": "Валовая маржа",
        "ebitda": "EBITDA маржа",
        "operations": "Операционная деятельность",
        "supplyChain": "Цепочка поставок",
        "criticalDeps": "Критические зависимости",
        "team": "Команда",
        "headcount": "Численность",
        "monthlyPayroll": "ФОТ / мес.",
        "annualPayroll": "ФОТ / год",
        "milestones": "Дорожная карта",
        "first90": "Первые 90 дней",
        "first6": "За 6 месяцев",
        "first12": "За 12 месяцев",
        "risks": "Риски",
        "kpis": "Ключевые показатели (KPI)",
        "products": "Рекомендуемые продукты НБУ",
        "nextSteps": "Что сделать перед визитом в банк",
        "creditScore": "Кредитный скоринг",
        "credit": {
            "verdicts": {
                "high": "Высокая кредитоспособность",
                "medium": "Средняя кредитоспособность",
                "low": "Низкая кредитоспособность",
            },
            "ratioNames": {
                "operatingMargin": "Операционная маржа",
                "netMargin": "Чистая маржа",
                "costToRevenue": "Расходы / выручка",
                "projectEquity": "Доля собственных в проекте",
                "loanToRevenue": "Кредит / годовая выручка",
                "dscr": "Покрытие платежей (DSCR)",
                "paybackYears": "Срок окупаемости",
                "businessAge": "Возраст бизнеса",
            },
            "methodology": (
                "Оценка рассчитана автоматически по 8 показателям из вашей анкеты в "
                "5 группах: прибыльность (операционная и чистая маржа), эффективность "
                "(расходы/выручка), структура капитала (доля собственных, кредит/выручка), "
                "обслуживание долга (DSCR, срок окупаемости), история бизнеса (возраст). "
                "Каждый показатель оценивается 0/1/2 балла, максимум — 16 баллов. "
                "Бэнды: ≥75% — высокая, ≥45% — средняя, <45% — низкая. "
                "Это самодиагностика на основе проектных данных, не решение банка."
            ),
        },
        "months": "мес.",
        "footer": "NBU AI Hub — Генератор бизнес-плана",
        "disclaimer": {
            "title": "Самодиагностика, не решение банка",
            "body": (
                "Этот документ и оценка кредитоспособности рассчитаны "
                "автоматически на основе данных анкеты. Это не профессиональная "
                "оценка кредитного аналитика и не итоговое решение банка. "
                "Используйте его, чтобы лучше понять свои шансы на одобрение и "
                "подготовиться к встрече в отделении НБУ — реальные условия и "
                "одобрение определяются на основании полной заявки и проверки документов."
            ),
        },
    },
    "uz": {
        "title": "Biznes-reja",
        "feasibility": "Loyiha amalga oshirish reytingi",
        "verdicts": {"high": "Yuqori", "medium": "Oʻrtacha", "low": "Past"},
        "executive": "Qisqacha xulosa",
        "market": "Bozor konteksti",
        "metrics": "Moliyaviy koʻrsatkichlar",
        "monthlyRevenue": "Oylik daromad",
        "monthlyCosts": "Oylik xarajat",
        "monthlyProfit": "Oylik foyda",
        "breakeven": "Tenglik nuqtasi",
        "grossMargin": "Yalpi marja",
        "ebitda": "EBITDA marja",
        "operations": "Ishlab chiqarish jarayoni",
        "supplyChain": "Yetkazib berish",
        "criticalDeps": "Asosiy bogʻliqliklar",
        "team": "Jamoa",
        "headcount": "Xodimlar soni",
        "monthlyPayroll": "Oylik ish haqi",
        "annualPayroll": "Yillik ish haqi",
        "milestones": "Reja qadamlari",
        "first90": "Birinchi 90 kun",
        "first6": "6 oy ichida",
        "first12": "12 oy ichida",
        "risks": "Xatarlar",
        "kpis": "Asosiy koʻrsatkichlar (KPI)",
        "products": "Tavsiya etilgan NBU kreditlari",
        "nextSteps": "Keyingi qadamlar — bankka tayyor borish uchun",
        "creditScore": "Kredit skoring",
        "credit": {
            "verdicts": {
                "high": "Yuqori kreditga layoqatlilik",
                "medium": "Oʻrtacha kreditga layoqatlilik",
                "low": "Past kreditga layoqatlilik",
            },
            "ratioNames": {
                "operatingMargin": "Operatsion marja",
                "netMargin": "Sof marja",
                "costToRevenue": "Xarajat / daromad",
                "projectEquity": "Loyihada oʻz ulush",
                "loanToRevenue": "Kredit / yillik daromad",
                "dscr": "Toʻlovlar qoplanishi (DSCR)",
                "paybackYears": "Qaytarilish muddati",
                "businessAge": "Biznes yoshi",
            },
            "methodology": (
                "Baho anketadagi 8 ta koʻrsatkich asosida 5 guruhda avtomatik "
                "hisoblangan: foydalilik (operatsion va sof marja), samaradorlik "
                "(xarajat/daromad), kapital tarkibi (oʻz ulush, kredit/daromad), "
                "qarz xizmati (DSCR, qaytarilish muddati), biznes tarixi (yosh). "
                "Har bir koʻrsatkich 0/1/2 ball bilan baholanadi, maksimum — 16 ball. "
                "Bandlar: ≥75% — yuqori, ≥45% — oʻrtacha, <45% — past. "
                "Bu loyiha maʼlumotlari boʻyicha oʻz-oʻzini diagnostika, bank qarori emas."
            ),
        },
        "months": "oy",
        "footer": "NBU AI Hub — Biznes-reja generatori",
        "disclaimer": {
            "title": "Oʻz-oʻzini diagnostika, bank qarori emas",
            "body": (
                "Ushbu hujjat va kreditga layoqatlilik bahosi anketa "
                "maʼlumotlari asosida avtomatik hisoblangan. Bu professional "
                "kredit tahlilchisining bahosi yoki bankning yakuniy qarori emas. "
                "Maʼqullanish ehtimolingizni yaxshiroq tushunish va NBU boʻlimida "
                "uchrashuvga tayyorlanish uchun foydalaning — haqiqiy shartlar "
                "va maʼqullash toʻliq ariza va hujjatlarni tekshirish asosida belgilanadi."
            ),
        },
    },
}


# ---------- helpers ----------

def _fmt(n: Any) -> str:
    try:
        return f"{int(n or 0):,}".replace(",", " ")
    except (TypeError, ValueError):
        return "0"


def _fmt_pct(n: Any) -> str:
    try:
        return f"{float(n or 0):.1f}%"
    except (TypeError, ValueError):
        return "0.0%"


# Simple "primary" colour for headings — matches the platform navy.
_PRIMARY = RGBColor(0x00, 0x3D, 0x7C)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _PRIMARY
        run.font.name = "Calibri"


def _para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Two-column key/value table for metrics blocks."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v
        for cell in (c0, c1):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def _bullets(doc: Document, items: list[str]) -> None:
    for it in items or []:
        if it:
            doc.add_paragraph(str(it), style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for it in items or []:
        if it:
            doc.add_paragraph(str(it), style="List Number")


# Verdict labels for the v2 shape; we render in Russian regardless of `lang`
# because DOCX is the formal export and the bank reads RU. (Frontend handles
# UZ display.)
_V2_VERDICTS = {
    "high": "Высокая кредитоспособность",
    "medium": "Средняя кредитоспособность",
    "low": "Низкая кредитоспособность",
    "needs_rework": "Бизнес-план требует доработки",
}

_V2_CRITERIA_LABELS = {
    "dscrSteadyState": "Покрытие платежей",
    "equityShare": "Структура капитала",
    "profitability": "Прибыльность",
    "realism": "Реалистичность",
    "resilience": "Устойчивость",
}

_V2_CRITERIA_ORDER = [
    "dscrSteadyState", "equityShare", "profitability", "realism", "resilience",
]

_V2_METHODOLOGY = (
    "Оценка рассчитана автоматически по 5 критериям, каждый 0-20 баллов "
    "(линейная интерполяция, без бинов): покрытие платежей в стационарном "
    "режиме (DSCR), структура капитала (доля собственных), прибыльность "
    "(EBITDA маржа vs медиана отрасли), реалистичность (число предупреждений "
    "из 4 проверок), устойчивость (DSCR в стресс-сценарии −20%/+10%). "
    "Итог 0-100. Бэнды: ≥75 — высокая, ≥50 — средняя, ≥25 — низкая, <25 — "
    "требует доработки. Это самодиагностика на основе анкеты, не решение банка."
)


def _render_credit_score_v2(doc: Document, score: dict[str, Any], t: dict) -> None:
    _heading(doc, t["creditScore"], 1)

    verdict_label = _V2_VERDICTS.get(score.get("verdict", "medium"), "")
    _para(doc, f"{verdict_label} — {score.get('total', 0)} / 100", bold=True)
    if score.get("summary"):
        _para(doc, score["summary"])

    industry = score.get("industry") or {}
    if industry.get("label"):
        _para(
            doc,
            f"Отрасль (классификация): {industry['label']} — медиана EBITDA маржи "
            f"{industry.get('ebitdaMarginMedian', 0)}%",
            size=10,
        )

    doc.add_paragraph()
    _para(doc, _V2_METHODOLOGY, size=10)

    # Per-criterion table — Criterion | Value | Points | Scale
    criteria = score.get("criteria") or {}
    rows_data = [
        (
            _V2_CRITERIA_LABELS.get(k, k),
            f"{criteria[k].get('value', '')}{criteria[k].get('unit', '')}",
            f"{criteria[k].get('points', 0)} / 20",
            criteria[k].get("scaleHint", ""),
        )
        for k in _V2_CRITERIA_ORDER
        if k in criteria
    ]
    if rows_data:
        doc.add_paragraph()
        table = doc.add_table(rows=len(rows_data) + 1, cols=4)
        table.style = "Light Grid Accent 1"
        h = table.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = (
            "Критерий", "Значение", "Баллы", "Шкала"
        )
        for i, (name, value, pts, scale) in enumerate(rows_data, start=1):
            row = table.rows[i].cells
            row[0].text = name
            row[1].text = value
            row[2].text = pts
            row[3].text = scale

    # Plausibility warnings
    flags = (criteria.get("realism") or {}).get("flags") or []
    if flags:
        doc.add_paragraph()
        _heading(doc, "Предупреждения по реалистичности", 2)
        for flag in flags:
            _para(doc, f"• {flag.get('message_ru', '')}", size=10)

    # Stress-test summary
    stress = score.get("stress") or {}
    if stress:
        doc.add_paragraph()
        _heading(doc, "Стресс-тест", 2)
        _para(
            doc,
            f"При снижении выручки на {int((1 - stress.get('revenueFactor', 0.8)) * 100)}% "
            f"и росте расходов на {int((stress.get('costFactor', 1.1) - 1) * 100)}%: "
            f"DSCR = {stress.get('dscr', 0):.2f}x",
            bold=True, size=10,
        )
        if stress.get("warning"):
            _para(
                doc,
                "⚠ В пессимистичном сценарии бизнес не покрывает платёж по кредиту.",
                size=10,
            )
        else:
            _para(
                doc,
                "Бизнес остаётся способным обслуживать кредит в пессимистичном сценарии.",
                size=10,
            )


def _render_credit_score_v1(doc: Document, credit_score: dict, t: dict) -> None:
    """Legacy v1 renderer — kept so DOCX downloads for old plans still work."""
    _heading(doc, t["creditScore"], 1)
    verdict = credit_score.get("verdict", "medium")
    verdict_label = t["credit"]["verdicts"].get(verdict, "")
    _para(
        doc,
        f"{verdict_label} — {credit_score.get('points', 0)}/{credit_score.get('maxPoints', 0)} "
        f"({credit_score.get('percent', 0)}%)",
        bold=True,
    )
    if credit_score.get("summary"):
        _para(doc, credit_score["summary"])

    methodology = t["credit"].get("methodology")
    if methodology:
        doc.add_paragraph()
        _para(doc, methodology, size=10)

    ratios = credit_score.get("ratios") or {}
    if ratios:
        doc.add_paragraph()
        table = doc.add_table(rows=len(ratios) + 1, cols=3)
        table.style = "Light Grid Accent 1"
        h = table.rows[0].cells
        h[0].text, h[1].text, h[2].text = "Показатель", "Значение", "Норма"
        ratio_names = t["credit"]["ratioNames"]
        for i, (key, info) in enumerate(ratios.items(), start=1):
            row = table.rows[i].cells
            row[0].text = ratio_names.get(key, key)
            row[1].text = f"{info.get('value', '')}{info.get('unit', '')}"
            row[2].text = str(info.get("benchmark") or "")


# ---------- builder ----------

def build_docx(
    *,
    plan: dict[str, Any],
    inputs: dict[str, Any],
    credit_score: dict[str, Any] | None,
    credit_score_v2: dict[str, Any] | None = None,
    lang: str = "ru",
) -> bytes:
    """Build a .docx and return raw bytes.

    If `credit_score_v2` is provided we render the v2 section (5 criteria,
    0-100, stress test, plausibility flags). Otherwise we fall back to the
    legacy v1 table — preserves DOCX downloads for plans created before v2
    rolled out.
    """
    t = L.get(lang) or L["ru"]
    doc = Document()

    # Title block
    org = (inputs or {}).get("organization") or {}
    org_name = org.get("name") or t["title"]
    title = doc.add_heading(t["title"], level=0)
    for r in title.runs:
        r.font.color.rgb = _PRIMARY
    sub = doc.add_paragraph()
    sub_run = sub.add_run(org_name)
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True

    meta_bits = []
    if org.get("mainActivity"):
        meta_bits.append(org["mainActivity"])
    if org.get("address"):
        meta_bits.append(org["address"])
    if org.get("inn"):
        meta_bits.append(f"ИНН: {org['inn']}")
    if meta_bits:
        m = doc.add_paragraph(" · ".join(meta_bits))
        for r in m.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Top disclaimer — placed right after the title block so readers see
    # "this is a self-diagnostic, not a bank decision" before any numbers.
    disclaimer = t.get("disclaimer")
    if disclaimer:
        doc.add_paragraph()
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(disclaimer["title"])
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        body_p = doc.add_paragraph()
        body_run = body_p.add_run(disclaimer["body"])
        body_run.font.size = Pt(10)
        body_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        doc.add_paragraph()

    # Verdict
    verdict = plan.get("feasibilityVerdict", "medium")
    score = plan.get("feasibilityScore", 0)
    _heading(doc, f"{t['feasibility']}: {t['verdicts'].get(verdict, '')} — {score}/100", 1)
    if plan.get("summary"):
        _para(doc, plan["summary"])

    # Executive summary
    if plan.get("executiveSummary"):
        _heading(doc, t["executive"], 1)
        _para(doc, plan["executiveSummary"])

    # Market context
    if plan.get("marketContext"):
        _heading(doc, t["market"], 1)
        _para(doc, plan["marketContext"])

    # Metrics block
    f = plan.get("financials") or {}
    c = f.get("monthlyCosts") or {}
    metric_rows = [
        (t["monthlyRevenue"], f"{_fmt(f.get('monthlyRevenue'))} UZS"),
        (t["monthlyCosts"], f"{_fmt(c.get('total'))} UZS"),
        (t["monthlyProfit"], f"{_fmt(f.get('monthlyProfit'))} UZS"),
        (t["breakeven"], f"{f.get('breakevenMonths', '—')} {t['months']}"),
        (t["grossMargin"], _fmt_pct(f.get("grossMarginPct"))),
        (t["ebitda"], _fmt_pct(f.get("ebitdaMarginPct"))),
    ]
    _heading(doc, t["metrics"], 1)
    _kv_table(doc, metric_rows)

    if f.get("assessment"):
        doc.add_paragraph()
        _para(doc, f["assessment"])

    # Operations
    ops = plan.get("operations") or {}
    if any(ops.values()):
        _heading(doc, t["operations"], 1)
        if ops.get("processFlow"):
            _bullets(doc, ops["processFlow"])
        if ops.get("supplyChain"):
            _para(doc, f"{t['supplyChain']}: {ops['supplyChain']}")
        if ops.get("criticalDependencies"):
            _para(doc, t["criticalDeps"] + ":", bold=True)
            _bullets(doc, ops["criticalDependencies"])

    # Team
    team = plan.get("team") or {}
    if team:
        _heading(doc, t["team"], 1)
        _kv_table(doc, [
            (t["headcount"], str(team.get("totalHeadcount", 0))),
            (t["monthlyPayroll"], f"{_fmt(team.get('monthlyPayroll'))} UZS"),
            (t["annualPayroll"], f"{_fmt(team.get('annualPayroll'))} UZS"),
        ])
        if team.get("assessment"):
            _para(doc, team["assessment"])

    # Milestones
    ms = plan.get("milestones") or {}
    if any(ms.values()):
        _heading(doc, t["milestones"], 1)
        for key, label in (("first90Days", t["first90"]),
                           ("first6Months", t["first6"]),
                           ("first12Months", t["first12"])):
            if ms.get(key):
                _heading(doc, label, 2)
                _bullets(doc, ms[key])

    # Risks
    risks = plan.get("risks") or []
    if risks:
        _heading(doc, t["risks"], 1)
        for r in risks:
            line = f"[{r.get('severity', '').upper()}] {r.get('description', '')}"
            _para(doc, line, bold=True)
            if r.get("mitigation"):
                _para(doc, f"→ {r['mitigation']}")

    # KPIs
    kpis = plan.get("kpis") or []
    if kpis:
        _heading(doc, t["kpis"], 1)
        table = doc.add_table(rows=len(kpis) + 1, cols=3)
        table.style = "Light Grid Accent 1"
        h = table.rows[0].cells
        h[0].text, h[1].text, h[2].text = "Показатель", "Цель", "Периодичность"
        for i, k in enumerate(kpis, start=1):
            row = table.rows[i].cells
            row[0].text = str(k.get("name") or "")
            row[1].text = str(k.get("target") or "")
            row[2].text = str(k.get("frequency") or "")

    # Recommended NBU products
    rec = plan.get("recommendedProducts") or []
    if rec:
        _heading(doc, t["products"], 1)
        for p in rec:
            _heading(doc, str(p.get("name") or ""), 2)
            _kv_table(doc, [
                ("Ставка", str(p.get("rate") or "")),
                ("Срок", str(p.get("term") or "")),
                ("Сумма", str(p.get("amount") or "")),
                ("Цель", str(p.get("purpose") or "")),
            ])
            if p.get("rationale"):
                doc.add_paragraph()
                _para(doc, str(p["rationale"]))

    # Action steps
    steps = plan.get("actionableNextSteps") or []
    if steps:
        _heading(doc, t["nextSteps"], 1)
        _numbered(doc, steps)

    # Credit scoring — v2 first (5 criteria, 0-100), v1 fallback
    if credit_score_v2:
        _render_credit_score_v2(doc, credit_score_v2, t)
    elif credit_score:
        _render_credit_score_v1(doc, credit_score, t)

    # Footer note
    doc.add_paragraph()
    f_para = doc.add_paragraph()
    f_run = f_para.add_run(t["footer"])
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
